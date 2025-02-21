import re
import json
import spacy
import glob
from collections import defaultdict
from docx import Document
import os
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,  # Set to DEBUG for more detailed logs
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("extraction.log"),
        logging.StreamHandler()
    ]
)

# Initialize spaCy NLP model
# Ensure that the "en_core_web_sm" model is installed. If not, install it using:
# python -m spacy download en_core_web_sm
try:
    nlp = spacy.load("en_core_web_sm")
    logging.info("spaCy model 'en_core_web_sm' loaded successfully.")
except Exception as e:
    logging.error(f"Failed to load spaCy model: {e}")
    raise e

# Define Penetration Testing Phases
PHASES = [
    "Reconnaissance",
    "Scanning",
    "Enumeration",
    "Exploitation",
    "Post-Exploitation",
    "Reporting",
    "Privilege Escalation",
    "Maintaining Access",
    "Covering Tracks",
    "Password Attacks",
    "Wireless Attacks",
    "Web Application Testing",
    "Sniffing and Spoofing",
    "Social Engineering",
    "Reverse Engineering",
    "Reporting Tools",
    "Miscellaneous Tools"
]

# ----------------------------------------------------------------------------
# Updated: Added several new CLI and Web tools in the lists below
# ----------------------------------------------------------------------------

# Define Command-Line Interface (CLI) Tools
CLI_TOOLS = [
    # Reconnaissance
    "Nmap", "theHarvester", "Recon-ng", "Dnsenum", "Amass", "Sublist3r",
    "Sherlock", "Censys", "SpiderFoot", "Netdiscover", "dnsrecon", "massdns",
    "aquatone",

    # Scanning and Enumeration
    "Masscan", "Gobuster", "SSLScan", "Enum4linux", "SNMPwalk",
    "Ldapsearch", "Fierce", "Netcat", "Zmap", "Hping3", "dirsearch", "DirBuster", "nikto",

    # Vulnerability Analysis
    "SQLmap", "Wapiti", "Arachni", "Skipfish", "Retire.js", "RIPS",
    "Vuls", "Nessus", "OpenVAS",

    # Exploitation
    "Metasploit Framework", "Canvas", "Core Impact",
    "SET (Social-Engineer Toolkit)", "PowerSploit", "EvilAP",
    "Responder", "Empire", "MSFvenom",

    # Post-Exploitation
    "PowerSploit", "Empire", "Mimikatz", "Responder", "Veil-Evasion",
    "Meterpreter", "PrivescCheck", "LaZagne", "Invoke-Obfuscation",
    "NetRipper", "Keylogger",

    # Password Attacks
    "John the Ripper", "Hashcat", "Hydra", "Medusa", "Crunch", "CeWL",
    "Patator", "RainbowCrack", "THC-Password-Cracker", "Ophcrack",
    "Cain", "Aircrack-ng", "WPA/WPA2 Crackers",

    # Wireless Attacks
    "Reaver", "Wifite", "Kismet", "Bettercap", "MDK3", "Cowpatty",
    "Airgeddon", "Wifiphisher",

    # Web Application Testing
    "W3af", "XSStrike", "Intruder", "IronWASP",

    # Sniffing and Spoofing
    "Bettercap", "Ettercap", "dsniff", "tcpdump",
    "MITMf (Man-In-The-Middle Framework)", "SSLstrip", "ARPSpoof",
    "MITMproxy", "Scapy", "Netcut",

    # Social Engineering
    "King Phisher", "Phishery", "Evilginx", "Credential Harvester",

    # Reverse Engineering
    "Radare2", "apktool", "Dex2jar", "Jadx", "Frida", "Capstone",

    # Reporting Tools
    "Dradis", "MagicTree", "Faraday", "Serpico", "CaseFile",
    "KeepNote", "Seas0nPass", "Pico", "Vega Report",

    # Miscellaneous Tools
    "Snort", "Cuckoo Sandbox", "Volatility",
    "Docker", "Vagrant", "VirtualBox", "VMware", "GPG", "Tor",
    "Proxychains", "Netcat", "Terminator", "tmux",
    "Autopsy", "Binwalk", "ExifTool", "Foremost",
    "Hash-identifier", "The Sleuth Kit", "Binjitsu",
    "Mitmproxy", "PowerShell Empire", "FuzzDB",
    "Sn1per", "OWASP Dependency-Check", "SecLists", "SecEdit",
    "Veil", "Wfuzz", "XSSer", "Yersinia", "Hashcat-utils",
    "Peepdf", "SecuriScan", "Tiger", "Unicornscan", "WFuzz",

    # Newly added / commonly used
    "Droopescan", "Sqlninja", "Dnsmap", "Ffuf", "Ghdb"
]

# Define Web-Based Tools
WEB_TOOLS = [
    # Reconnaissance
    "Maltego", "OSINT Framework", "FOCA", "WhatWeb", "Censys",

    # Vulnerability Analysis
    "Burp Suite", "Vega",

    # Exploitation
    "BeEF (Browser Exploitation Framework)", "Armitage",

    # Web Application Testing
    "OWASP ZAP (Zed Attack Proxy)", "Burp Suite", "IronWASP",

    # Social Engineering
    "BeEF", "Gophish", "HiddenEye",

    # Reporting Tools
    "Metasploit Pro", "Burp Suite Professional",

    # Reverse Engineering
    "Ghidra", "IDA Pro", "Binary Ninja", "Cutter",
    "x64dbg", "Immunity Debugger",

    # Miscellaneous Tools
    "Metasploit Pro", "Cobalt Strike", "Autopsy",
    "ExploitDB",

    # Newly added
    "Acunetix", "Nexpose"
]

# Define non-command phrases to exclude
NON_COMMAND_PHRASES = [
    'scan report',
    'host is up',
    'starting at',
    'report for',
    'done:',
    'scanned in',
    'tips:',
    'tips to',
    'imagine the following scenario',
    'environment',
    'dashboard',
    'warning',
    'comment',
    'username',  # might remove if needed
    'password',  # might remove if needed
    'secret',
    'explain',
    'overview',
    'interface',
    'discovery',
    'enumerating',
    'snapshot',
    'attacking',
    'inject',
    'payload'
]

def load_docx(file_path):
    """
    Load and parse the DOCX file.
    Returns a list of paragraphs with their text and styles.
    """
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                'text': para.text.strip(),
                'style': para.style.name
            })
        return paragraphs
    except Exception as e:
        logging.error(f"Error loading DOCX file {file_path}: {e}")
        return []

def list_unique_styles(paragraphs, file_name):
    """
    Lists all unique styles present in the DOCX file.
    Useful for debugging and ensuring correct style detection.
    """
    styles = set()
    for para in paragraphs:
        styles.add(para['style'])
    logging.info(f"\nUnique Styles Found in {file_name}:")
    for style in sorted(styles):
        logging.info(f"- {style}")

def split_into_sections(paragraphs):
    """
    Split paragraphs into sections based on phase headings.
    """
    sections = {}
    current_phase = None

    for para in paragraphs:
        text = para['text']
        style = para['style']

        if not text:
            continue  # Skip empty paragraphs

        # Detect if the paragraph is a heading matching any phase
        for phase in PHASES:
            if re.search(r'\b' + re.escape(phase) + r'\b', text, re.IGNORECASE):
                current_phase = phase
                if current_phase not in sections:
                    sections[current_phase] = []
                logging.info(f"Detected phase: {current_phase}")
                break
        else:
            if current_phase:
                sections[current_phase].append(para)

    return sections

def extract_commands_from_tables(doc):
    """
    Extract commands from tables within the DOCX file.
    Returns a list of (command, fields) tuples.
    """
    commands = []
    ip_pattern = r'(\b(?:\d{1,3}\.){3}\d{1,3}(?:/\d{1,2})?\b)|(\b(?:[A-Fa-f0-9]{1,4}:){7}[A-Fa-f0-9]{1,4}\b)'

    for table in doc.tables:
        for row in table.rows:
            # If there's a header row you want to skip, you can do so here
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    # Updated: detect commands that might start with a path or alias
                    combined_tools = CLI_TOOLS + WEB_TOOLS

                    # This pattern looks for optional path or alias like /usr/bin/, ./, etc.
                    # e.g.: /usr/bin/nmap -sV or ./nmap -sV
                    cmd_pattern = (
                        r'^((?:\.?/[\w/]+)?('
                        + '|'.join([re.escape(tool) for tool in combined_tools])
                        + r'))\b\s+[-/]'
                    )

                    if re.match(cmd_pattern, text, re.IGNORECASE):
                        # Exclude non-command phrases
                        if not any(phrase in text.lower() for phrase in NON_COMMAND_PHRASES):
                            # Detect and replace IPs
                            matches = re.findall(ip_pattern, text)
                            fields = []
                            if matches:
                                text = re.sub(ip_pattern, '{ip}', text)
                                fields.append('ip')
                            # Clean the command
                            cleaned_cmd, cmd_fields = clean_command(text)
                            if cleaned_cmd:
                                fields.extend(cmd_fields)
                                commands.append((cleaned_cmd, fields))
                                logging.info(f"Extracted Table Command: {cleaned_cmd} | Fields: {fields}")
    return commands

def extract_commands_from_text(text):
    """
    Extract commands embedded within regular text.
    Returns a list of (command, fields) tuples.
    """
    # Regex to find commands starting with optional paths or recognized tool names
    combined_tools = CLI_TOOLS + WEB_TOOLS
    cmd_pattern = (
        r'((?:\.?/[\w/]+)?('
        + '|'.join([re.escape(tool) for tool in combined_tools])
        + r'))\b\s+[-/][^\s]+(?:\s+[^\s]+)*'
    )

    matches = re.finditer(cmd_pattern, text, re.IGNORECASE)
    full_commands = []
    for match in matches:
        # Extract the full command line
        cmd_text = match.group()
        # Heuristic to assume command ends at period or line break, etc.
        end_match = re.search(r'[.\n]', text[text.find(cmd_text):])
        if end_match:
            end = text.find(cmd_text) + end_match.start()
            cmd = text[text.find(cmd_text):end]
        else:
            cmd = cmd_text
        cmd = cmd.strip()
        # Clean the command and get required fields
        cleaned_cmd, fields = clean_command(cmd)
        if cleaned_cmd:
            full_commands.append((cleaned_cmd, fields))
    return full_commands

def clean_command(cmd):
    """
    Clean and validate extracted commands while retaining placeholders.
    Detects IP addresses and replaces them with {ip}.
    """
    # ----------------------------------------------------------------------------
    # UPDATED: Added some extra placeholders: <username>, <password>, <path>, <domain>
    # ----------------------------------------------------------------------------
    placeholder_mapping = {
        '<Target URL>': '{target_url}',
        '[IP Address]': '{ip}',
        '[port number]': '{port}',
        '[Hash Type]': '{hash_type}',
        '[Output File Path]': '{output_path}',
        '[Hashes File Path]': '{hashes_path}',
        '[Dictionary File Path]': '{dict_path}',
        '[URL]': '{url}',
        '[db]': '{database}',
        '[db name]': '{db_name}',
        '[table name]': '{table_name}',
        '[NSE script name]': '{nse_script}',

        # <-- Newly added placeholders
        '<username>': '{username}',
        '<password>': '{password}',
        '<path>': '{path}',
        '<domain>': '{domain}',
    }

    # Replace predefined placeholders with standardized tokens
    for placeholder, token in placeholder_mapping.items():
        if placeholder in cmd:
            cmd = cmd.replace(placeholder, token).strip()

    # Detect and replace IP addresses with {ip}
    ip_pattern = r'(\b(?:\d{1,3}\.){3}\d{1,3}(?:/\d{1,2})?\b)|(\b(?:[A-Fa-f0-9]{1,4}:){7}[A-Fa-f0-9]{1,4}\b)'
    matches = re.findall(ip_pattern, cmd)
    fields = []

    if matches:
        cmd = re.sub(ip_pattern, '{ip}', cmd)
        fields.append('ip')

    # Remove trailing or leading special characters
    cmd = cmd.strip('[]<>')

    # Remove variable assignments or incomplete commands
    # e.g., "TARGET_IP=1.2.3.4" or "port="
    if '=' in cmd:
        return None, []

    # Exclude commands that are too short or likely non-commands
    if len(cmd.split()) < 2:
        return None, []

    return cmd, fields

def deduplicate_commands(commands):
    """
    Remove duplicate and empty commands from the list.
    Returns a list of unique command dictionaries with 'command' and 'fields'.
    """
    unique_commands = {}
    for cmd, fields in commands:
        if cmd and cmd not in unique_commands:
            unique_commands[cmd] = set(fields)
        elif cmd:
            unique_commands[cmd].update(fields)
    
    deduped = []
    for cmd, fields in unique_commands.items():
        deduped.append({
            'command': cmd,
            'fields': sorted(list(fields))
        })
    return deduped

def extract_tools_and_commands(section_paragraphs, doc):
    """
    Extract tools and commands from a section's paragraphs and tables.
    Returns a tuple of (cli_tools_found, web_tools_found, commands_found)
    """
    cli_tools_found = set()
    web_tools_found = set()
    commands_found = []

    for para in section_paragraphs:
        text = para['text']
        style = para['style']

        if not text:
            continue

        # Identify CLI tools in the paragraph
        for tool in CLI_TOOLS:
            if re.search(r'\b' + re.escape(tool) + r'\b', text, re.IGNORECASE):
                cli_tools_found.add(tool)

        # Identify Web-based tools in the paragraph
        for tool in WEB_TOOLS:
            if re.search(r'\b' + re.escape(tool) + r'\b', text, re.IGNORECASE):
                web_tools_found.add(tool)

        # Identify commands based on style or pattern
        is_command = False
        required_fields = []

        # Check if the paragraph style indicates code or preformatted text
        command_styles = ['Code', 'Preformatted', 'Code Block', 'Consolas', 'Courier New', 'Monospace']
        if style.lower() in [s.lower() for s in command_styles]:
            is_command = True
        else:
            # Additionally, detect commands with optional path prefix
            combined_tools = CLI_TOOLS + WEB_TOOLS
            cmd_pattern = (
                r'^((?:\.?/[\w/]+)?('
                + '|'.join([re.escape(tool) for tool in combined_tools])
                + r'))\b\s+[-/]'
            )
            if re.match(cmd_pattern, text, re.IGNORECASE):
                if not any(phrase in text.lower() for phrase in NON_COMMAND_PHRASES):
                    is_command = True

        if is_command:
            cleaned_cmd, fields = clean_command(text)
            if cleaned_cmd:
                commands_found.append({
                    'command': cleaned_cmd,
                    'fields': fields
                })
                logging.info(f"Extracted Command: {cleaned_cmd} | Fields: {fields}")
        else:
            # Extract inline commands within the paragraph
            inline_commands = extract_commands_from_text(text)
            for cmd, fields in inline_commands:
                if not any(phrase in cmd.lower() for phrase in NON_COMMAND_PHRASES):
                    commands_found.append({
                        'command': cmd,
                        'fields': fields
                    })
                    logging.info(f"Extracted Inline Command: {cmd} | Fields: {fields}")

    # Extract commands from tables in this section
    table_commands = extract_commands_from_tables(doc)
    for cmd, fields in table_commands:
        if not any(phrase in cmd.lower() for phrase in NON_COMMAND_PHRASES):
            commands_found.append({
                'command': cmd,
                'fields': fields
            })
            logging.info(f"Extracted Table Command: {cmd} | Fields: {fields}")

    # Deduplicate
    commands_found = deduplicate_commands([(c['command'], c['fields']) for c in commands_found])

    return cli_tools_found, web_tools_found, commands_found

# ----------------------------------------------------------------------------
# UPDATED: multi-line merging now looks for both "\" and "&&" as line continuations
# ----------------------------------------------------------------------------
def preprocess_paragraphs(paragraphs):
    """
    Merge consecutive command paragraphs that are part of the same multi-line command.
    Extends beyond just '\' to also include '&&' as a potential line continuation.
    """
    merged_paragraphs = []
    current_command = ""
    command_styles = ['Code', 'Preformatted', 'Code Block', 'Consolas', 'Courier New', 'Monospace']

    for para in paragraphs:
        text = para['text']
        style = para['style']

        if not text:
            continue

        # Check if the style is recognized as code-like
        if style.lower() in [s.lower() for s in command_styles]:
            # Look for line continuations
            if text.endswith('\\'):
                current_command += text[:-1].strip() + ' '
            elif text.endswith('&&'):
                # If line ends with '&&', keep accumulating
                current_command += text[:-2].strip() + ' && '
            else:
                # Final line in the multi-line command
                current_command += text
                merged_paragraphs.append({'text': current_command, 'style': style})
                current_command = ""
        else:
            # If we encounter a non-code paragraph, finalize any open command first
            if current_command:
                merged_paragraphs.append({'text': current_command, 'style': 'Code'})
                current_command = ""
            merged_paragraphs.append(para)

    # If something is still left in current_command
    if current_command:
        merged_paragraphs.append({'text': current_command, 'style': 'Code'})

    return merged_paragraphs

def build_knowledge_base(paragraphs, doc):
    """
    Build a knowledge base dictionary from the list of paragraphs.
    """
    knowledge_base = defaultdict(lambda: {'cli_tools': set(), 'web_tools': set(), 'commands': defaultdict(list)})
    sections = split_into_sections(paragraphs)

    logging.info(f"\nTotal sections detected: {len(sections)}")
    for phase, section_paragraphs in sections.items():
        cli_tools, web_tools, commands = extract_tools_and_commands(section_paragraphs, doc)

        # Add CLI Tools
        for tool in cli_tools:
            knowledge_base[phase]['cli_tools'].add(tool)

        # Add Web-Based Tools
        for tool in web_tools:
            knowledge_base[phase]['web_tools'].add(tool)

        for cmd in commands:
            # Determine which tool the command belongs to
            matched_tool = None
            combined_tools = CLI_TOOLS + WEB_TOOLS
            cmd_lower = cmd['command'].lower()
            for tool in combined_tools:
                # Check if the command starts with the tool name or path alias
                # e.g., '/usr/bin/nmap -sV' or 'nmap -sV'
                pattern_tool = r'^(\.?/[\w/]+)?' + re.escape(tool.lower()) + r'\b'
                if re.match(pattern_tool, cmd_lower):
                    matched_tool = tool
                    break

            if matched_tool:
                if matched_tool in CLI_TOOLS:
                    knowledge_base[phase]['commands'][matched_tool].append({
                        'command': cmd['command'],
                        'fields': cmd['fields'],
                        'type': 'CLI'
                    })
                elif matched_tool in WEB_TOOLS:
                    knowledge_base[phase]['commands'][matched_tool].append({
                        'command': cmd['command'],
                        'fields': cmd['fields'],
                        'type': 'Web'
                    })
            else:
                knowledge_base[phase]['commands']['Other'].append({
                    'command': cmd['command'],
                    'fields': cmd['fields'],
                    'type': 'Unknown'
                })

        # Debugging Output
        logging.info(f"\nPhase: {phase}")
        logging.info(f"CLI Tools Found: {cli_tools}")
        logging.info(f"Web Tools Found: {web_tools}")
        for tool, cmds in knowledge_base[phase]['commands'].items():
            logging.info(f"\n  Tool: {tool}")
            for c in cmds:
                logging.info(f"    - {c['command']} | Fields: {c['fields']} | Type: {c['type']}")

    # Convert sets to lists and ensure JSON serializable structure
    kb_serializable = {
        phase: {
            'cli_tools': sorted(list(details['cli_tools'])),
            'web_tools': sorted(list(details['web_tools'])),
            'commands': {tool: cmds for tool, cmds in details['commands'].items()}
        }
        for phase, details in knowledge_base.items()
    }

    return kb_serializable

def save_knowledge_base(kb, output_file):
    """
    Save the knowledge base to a JSON file.
    """
    try:
        # Print the knowledge base to verify its content
        logging.info("\nKnowledge Base to Save:")
        logging.info(json.dumps(kb, indent=4))

        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(kb, f, indent=4)
        logging.info(f"\nKnowledge base successfully saved to {output_file}")
    except Exception as e:
        logging.error(f"Error saving knowledge base: {e}")

def display_knowledge_base(kb):
    """
    Display commands for each phase and tool.
    """
    for phase, details in kb.items():
        logging.info(f"\nPhase: {phase}")
        logging.info(f"CLI Tools: {', '.join(details['cli_tools'])}")
        logging.info(f"Web Tools: {', '.join(details['web_tools'])}")
        for tool, commands in details['commands'].items():
            logging.info(f"\n  Tool: {tool}")
            for cmd in commands:
                logging.info(f"    - {cmd['command']} | Fields: {cmd['fields']} | Type: {cmd['type']}")

def main():
    """
    Main function to build the knowledge base.
    """
    # Directory containing DOCX files
    # Make this path dynamic or configurable based on the deployment environment
    # For cross-platform compatibility, use os.path.join and relative paths if possible
    docx_directory = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Docx')  # Assuming DOCX files are in 'Docx' folder within the project

    # Output JSON file
    output_kb_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'knowledge_base.json')

    # Find all DOCX files in the specified directory
    docx_files = glob.glob(os.path.join(docx_directory, '*.docx'))

    if not docx_files:
        logging.error(f"No DOCX files found in the directory: {docx_directory}")
        return

    # Initialize the aggregated knowledge base
    aggregated_kb = defaultdict(lambda: {'cli_tools': set(), 'web_tools': set(), 'commands': defaultdict(list)})

    # Process each DOCX file
    for file_path in docx_files:
        file_name = os.path.basename(file_path)
        logging.info(f"\nProcessing file: {file_name}")

        # Load and parse DOCX file
        paragraphs = load_docx(file_path)

        if not paragraphs:
            logging.warning(f"Skipping file due to loading issues: {file_name}")
            continue

        # List unique styles for debugging
        list_unique_styles(paragraphs, file_name)

        # Preprocess paragraphs to merge multi-line commands
        paragraphs = preprocess_paragraphs(paragraphs)

        # Load the document again for table extraction
        doc = Document(file_path)

        # Build knowledge base from the current file
        kb = build_knowledge_base(paragraphs, doc)

        # Aggregate the current file's knowledge base into the aggregated_kb
        for phase, details in kb.items():
            aggregated_kb[phase]['cli_tools'].update(details['cli_tools'])
            aggregated_kb[phase]['web_tools'].update(details['web_tools'])
            for tool, cmds in details['commands'].items():
                aggregated_kb[phase]['commands'][tool].extend(cmds)

    # Convert sets to lists and ensure JSON serializable structure
    kb_serializable = {
        phase: {
            'cli_tools': sorted(list(details['cli_tools'])),
            'web_tools': sorted(list(details['web_tools'])),
            'commands': {tool: cmds for tool, cmds in details['commands'].items()}
        }
        for phase, details in aggregated_kb.items()
    }

    # Save the aggregated knowledge base to a JSON file
    save_knowledge_base(kb_serializable, output_kb_file)

    # Optional: Display the final aggregated knowledge base
    display_knowledge_base(kb_serializable)

def start_extraction():
    """
    Function to initiate the extraction process.
    Designed to be called by the Flask backend.
    """
    try:
        main()
    except Exception as e:
        logging.error(f"An error occurred during extraction: {e}")
        raise e

if __name__ == "__main__":
    main()
